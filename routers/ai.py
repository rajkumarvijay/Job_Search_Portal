"""
AI-powered routes:
  GET  /api/v1/ai/jobs/search        — world-wide job search via HuggingFace
  POST /api/v1/ai/resume             — resume ATS analysis via HuggingFace
  POST /api/v1/ai/resume/full        — ATS analysis + AI job recommendations
"""

import io
import logging
from fastapi import APIRouter, Query, UploadFile, File, Form, HTTPException
from services.hf_llm_service import (
    search_jobs_ai,
    analyze_resume,
    get_resume_job_recommendations,
    generate_cover_letter,
    match_resume_to_job,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/ai", tags=["ai"])

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB


# ── AI job search ──────────────────────────────────────────────────────────────
@router.get("/jobs/search")
async def ai_job_search(
    q:              str = Query(..., min_length=1),
    location:       str = Query("worldwide"),
    results_wanted: int = Query(20, ge=5, le=40),
):
    try:
        jobs = await search_jobs_ai(q, location, results_wanted)
        return {"query": q, "location": location, "total": len(jobs), "source": "huggingface-ai", "jobs": jobs}
    except ValueError as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(f"AI job search error: {e}")
        raise HTTPException(status_code=500, detail="AI search failed")


# ── Resume ATS analysis (score only) ──────────────────────────────────────────
@router.post("/resume")
async def analyse_resume(
    file:        UploadFile = File(...),
    target_role: str        = Form(default=""),
):
    resume_text = await _read_resume(file)
    try:
        result = await analyze_resume(resume_text, target_role)
        return result
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(f"Resume analysis error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Resume analysis failed: {e}")


# ── Resume full analysis + job recommendations ─────────────────────────────────
@router.post("/resume/full")
async def analyse_resume_full(
    file:        UploadFile = File(...),
    target_role: str        = Form(default=""),
    job_count:   int        = Form(default=6),
):
    resume_text = await _read_resume(file)
    try:
        ats_result = await analyze_resume(resume_text, target_role)

        if ats_result.get("error"):
            return {**ats_result, "recommended_jobs": []}

        skills = ats_result.get("top_skills", [])
        roles  = ats_result.get("recommended_roles", [])
        level  = ats_result.get("experience_level", "Mid-Level")

        jobs = await get_resume_job_recommendations(
            skills=skills,
            roles=roles,
            level=level,
            target_role=target_role,
            count=max(3, min(job_count, 9)),
        )

        return {**ats_result, "recommended_jobs": jobs}

    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(f"Full resume analysis error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Resume analysis failed: {e}")


# ── Job Match Score ────────────────────────────────────────────────────────────
@router.post("/job-match")
async def ai_job_match(
    job_description: str        = Form(...),
    file:            UploadFile = File(default=None),
    resume_text:     str        = Form(default=""),
):
    """
    Score how well a resume matches a job description.
    Returns match score, matched skills, missing skills, strengths, learning recommendations.
    """
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description is required")

    extracted = ""
    if file and file.filename:
        extracted = await _read_resume(file)
    elif resume_text.strip():
        extracted = resume_text.strip()
    else:
        raise HTTPException(status_code=400, detail="Provide a resume file or paste your resume text")

    try:
        result = await match_resume_to_job(
            resume_text=extracted,
            job_description=job_description,
        )
        return result
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(f"Job match error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Job match analysis failed: {e}")


# ── Cover Letter Generator ─────────────────────────────────────────────────────
@router.post("/cover-letter")
async def ai_cover_letter(
    job_description:  str        = Form(...),
    candidate_name:   str        = Form(default=""),
    company_name:     str        = Form(default=""),
    job_title:        str        = Form(default=""),
    tone:             str        = Form(default="Professional"),
    file:             UploadFile = File(default=None),
    resume_text:      str        = Form(default=""),
):
    """
    Generate a tailored cover letter from resume + job description.
    Accepts either a resume file upload OR pasted resume_text.
    """
    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description is required")

    extracted = ""
    if file and file.filename:
        extracted = await _read_resume(file)
    elif resume_text.strip():
        extracted = resume_text.strip()
    else:
        raise HTTPException(status_code=400, detail="Provide a resume file or paste your resume text")

    try:
        result = await generate_cover_letter(
            resume_text=extracted,
            job_description=job_description,
            candidate_name=candidate_name,
            company_name=company_name,
            job_title=job_title,
            tone=tone,
        )
        return result
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=503, detail=str(e))
    except Exception as e:
        logger.error(f"Cover letter error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Cover letter generation failed: {e}")


# ── Helpers ────────────────────────────────────────────────────────────────────
async def _read_resume(file: UploadFile) -> str:
    ct   = file.content_type or ""
    name = (file.filename or "").lower()

    allowed_ext  = (".pdf", ".docx", ".txt")
    allowed_ct   = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
    }

    if ct not in allowed_ct and not any(name.endswith(e) for e in allowed_ext):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, or TXT files are accepted")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large (max 5 MB)")

    text = _extract_text(content, name, ct)
    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from the file. Make sure it is not a scanned/image-only PDF.",
        )
    return text


def _extract_text(content: bytes, filename: str, content_type: str) -> str:
    if filename.endswith(".pdf") or "pdf" in content_type:
        return _extract_pdf(content)
    if filename.endswith(".docx") or "wordprocessingml" in content_type:
        return _extract_docx(content)
    return content.decode("utf-8", errors="ignore")


def _extract_pdf(content: bytes) -> str:
    """
    Try pdfplumber first (preserves table/column layout better),
    fall back to pypdf if pdfplumber yields nothing.
    """
    text = _pdf_pdfplumber(content)
    if text.strip():
        return text
    logger.warning("[PDF] pdfplumber returned empty — falling back to pypdf")
    return _pdf_pypdf(content)


def _pdf_pdfplumber(content: bytes) -> str:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                # Also pull text from any tables on the page
                for table in page.extract_tables():
                    for row in table:
                        row_text = " | ".join(cell or "" for cell in row if cell)
                        if row_text.strip():
                            page_text += "\n" + row_text
                pages.append(page_text)
        return "\n".join(pages)
    except Exception as e:
        logger.warning(f"[PDF] pdfplumber error: {e}")
        return ""


def _pdf_pypdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        logger.warning(f"[PDF] pypdf error: {e}")
        return ""


def _extract_docx(content: bytes) -> str:
    """
    Use python-docx to extract paragraphs and table cells,
    preserving section structure (headings, bullets, tables).
    Falls back to raw XML parsing if python-docx is unavailable.
    """
    text = _docx_python_docx(content)
    if text.strip():
        return text
    logger.warning("[DOCX] python-docx returned empty — falling back to XML")
    return _docx_xml(content)


def _docx_python_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        lines = []

        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)

        return "\n".join(lines)
    except Exception as e:
        logger.warning(f"[DOCX] python-docx error: {e}")
        return ""


def _docx_xml(content: bytes) -> str:
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        texts = [node.text for node in tree.iter(f"{ns}t") if node.text]
        return " ".join(texts)
    except Exception as e:
        logger.warning(f"[DOCX] XML fallback error: {e}")
        return ""
